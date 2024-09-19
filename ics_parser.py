import streamlit as st
from ics import Calendar
import requests
import pandas as pd
from datetime import datetime, timedelta
import base64
from io import BytesIO
from docx import Document

# Password for the app
PASSWORD = "your_password_here"

# Default ICS link
DEFAULT_ICS_URL = "https://outlook.office365.com/owa/calendar/1cd1c906443845f3b6f75a99e0046625@hlw-stpeter.at/eb0bfd4af91541c186aca61ab066659016968059287048739671/calendar.ics"

def fetch_ics_data(ics_url):
    try:
        response = requests.get(ics_url)
        response.raise_for_status()  # Check if the request was successful
        return response.text
    except requests.exceptions.RequestException as e:
        st.error(f"Fehler beim Abrufen der ICS-Datei: {e}")
        return None

def parse_ics_data(calendar_data, start_date, end_date):
    calendar = Calendar(calendar_data)
    events_list = []
    
    for event in calendar.events:
        event_start = event.begin.datetime
        event_end = event.end.datetime if event.end else None

        # Make event_start and event_end naive if they are timezone-aware
        if event_start.tzinfo is not None:
            event_start = event_start.replace(tzinfo=None)
        if event_end and event_end.tzinfo is not None:
            event_end = event_end.replace(tzinfo=None)
        
        if event_start < start_date or (event_end and event_end > end_date):
            continue  # Skip events outside the specified range

        # Determine if the end date should be shown
        end_date_str = event_end.strftime("%d. %m. %Y") if event_end and (event_end - event_start) > timedelta(days=1) else ""
        
        # Only show time if it’s not 00:00
        time_str = event_start.strftime("%H:%M") if event_start.time() != datetime.min.time() else ""
        
        events_list.append({
            'Startdatum': event_start.strftime("%d. %m. %Y"),
            'Enddatum': end_date_str,
            'Uhrzeit': time_str,
            'Ereignis': event.name,
            'Startdatetime': event_start  # Add for sorting purposes
        })

    # Convert to DataFrame and sort by Startdatetime
    df = pd.DataFrame(events_list)
    df = df.sort_values(by='Startdatetime').reset_index(drop=True)
    
    return df

def display_table(df):
    st.dataframe(df[['Startdatum', 'Enddatum', 'Uhrzeit', 'Ereignis']])
    
    # Add buttons to download the table in different formats
    st.markdown(get_table_download_link(df, 'docx'), unsafe_allow_html=True)  # Word download as first option
    st.markdown(get_table_download_link(df, 'csv'), unsafe_allow_html=True)
    st.markdown(get_table_download_link(df, 'xlsx'), unsafe_allow_html=True)

def get_table_download_link(df, file_type):
    if file_type == 'csv':
        csv = df.to_csv(index=False)
        b64 = base64.b64encode(csv.encode()).decode()  # Convert to base64
        return f'<a href="data:file/csv;base64,{b64}" download="kalender_ereignisse.csv">CSV herunterladen</a>'
    elif file_type == 'xlsx':
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Ereignisse')
        b64 = base64.b64encode(output.getvalue()).decode()  # Convert to base64
        return f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="kalender_ereignisse.xlsx">Excel herunterladen</a>'
    elif file_type == 'docx':
        doc = Document()
        doc.add_heading('Kalenderereignisse', 0)
        table = doc.add_table(rows=1, cols=len(df.columns))
        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(df.columns):
            hdr_cells[i].text = column_name

        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

        output = BytesIO()
        doc.save(output)
        b64 = base64.b64encode(output.getvalue()).decode()  # Convert to base64
        return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="kalender_ereignisse.docx">Word herunterladen</a>'

def main():
    st.title("ICS-Kalenderparser")
    
    # Password authentication
    if 'authenticated' not in st.session_state or not st.session_state.authenticated:
        st.subheader("Bitte anmelden")
        password = st.text_input("Passwort", type="password")
        
        # Confirmation button for password
        if st.button("Anmelden"):
            if password == PASSWORD:
                st.session_state.authenticated = True
            else:
                st.error("Falsches Passwort")
        return

    # Main content of the app
    ics_url = st.text_input("Geben Sie den ICS-Kalenderlink ein:", DEFAULT_ICS_URL)
    
    # Confirmation button for ICS URL
    if st.button("Kalender laden"):
        # Default date range
        current_year = datetime.now().year
        start_date = datetime(current_year, 9, 1)
        end_date = datetime(current_year + 1, 9, 1)

        # User inputs for date range
        st.sidebar.header("Filterdatum")
        start_date = st.sidebar.date_input("Startdatum", start_date)
        end_date = st.sidebar.date_input("Enddatum", end_date)

        # Convert to datetime objects
        start_date = datetime.combine(start_date, datetime.min.time())
        end_date = datetime.combine(end_date, datetime.min.time())

        if ics_url:
            calendar_data = fetch_ics_data(ics_url)
            
            if calendar_data:
                df = parse_ics_data(calendar_data, start_date, end_date)
                display_table(df)

if __name__ == "__main__":
    main()
